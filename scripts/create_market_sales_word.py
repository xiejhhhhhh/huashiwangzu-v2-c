from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path("/Users/hekunhua/Documents/Hermes工作区/竞品分析/全肌模式招商电销面销话术_市场部训练版.md")
OUTPUT = Path("/Users/hekunhua/Documents/Hermes工作区/竞品分析/全肌模式招商电销面销话术_市场部培训版.docx")

FONT_CN = "Microsoft YaHei"
FONT_EN = "Aptos"
INK = "1D1D1F"
MUTED = "6E6B66"
ORANGE = "F15A3A"
GREEN = "28776D"
BLUE = "3778A8"
LINE = "D8D2C8"
PALE_ORANGE = "FCE4D6"
PALE_GREEN = "E2F0D9"
PALE_BLUE = "DDEBF7"
PALE_YELLOW = "FFF2CC"
PALE_GREY = "F2F2F2"


def set_run_font(run, size: float, color: str = INK, bold: bool = False, italic: bool = False):
    run.font.name = FONT_EN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)


def set_shading(element, fill: str):
    if hasattr(element, "_tc"):
        props = element._tc.get_or_add_tcPr()
    else:
        props = element._p.get_or_add_pPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: int = 5):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_left_border(paragraph, color: str = ORANGE, size: int = 18):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    return run


def add_inline_runs(paragraph, text: str, size: float = 13.5, color: str = INK):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, color, True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size - 0.5, ORANGE, True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, color)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(13.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    specs = {
        "Title": (29, INK, True),
        "Subtitle": (15, MUTED, False),
        "Heading 1": (21, ORANGE, True),
        "Heading 2": (17.5, INK, True),
        "Heading 3": (15.2, GREEN, True),
        "Heading 4": (14, BLUE, True),
    }
    for name, (size, color, bold) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT_EN
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    header = section.header
    p = header.paragraphs[0]
    r = p.add_run("全肌模式招商电销与面销话术｜市场部培训版")
    set_run_font(r, 8.5, MUTED)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第 ")
    set_run_font(r, 8.5, MUTED)
    r = add_field(p, "PAGE")
    set_run_font(r, 8.5, MUTED)
    r = p.add_run(" 页 / 共 ")
    set_run_font(r, 8.5, MUTED)
    r = add_field(p, "NUMPAGES")
    set_run_font(r, 8.5, MUTED)
    r = p.add_run(" 页")
    set_run_font(r, 8.5, MUTED)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(64)
    r = p.add_run("市场部招商转化训练手册")
    set_run_font(r, 13, ORANGE, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run("全肌模式招商")
    set_run_font(r, 29, INK, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("电销话术 + 面销话术")
    set_run_font(r, 24, GREEN, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("从线索筛选、门店诊断到异议处理与签约收口")
    set_run_font(r, 14.5, MUTED)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    for cell, fill in [(left, PALE_ORANGE), (right, PALE_GREEN)]:
        cell.width = Cm(7.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_shading(cell, fill)
        set_cell_border(cell, fill, 1)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run("电销\n筛选 · 诊断 · 预约")
    set_run_font(r, 14, INK, True)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run("面销\n盘店 · 讲解 · 成交")
    set_run_font(r, 14, INK, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(115)
    r = p.add_run("内部训练资料｜竞品模式转译｜非俏小喵正式招商政策")
    set_run_font(r, 10.5, MUTED)
    doc.add_page_break()


def add_quick_rules(doc: Document, source: str):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("使用说明")
    rules = [
        ("电销目标", "不在电话里讲完整方案，只成交一次有效诊断或面谈。", PALE_ORANGE),
        ("面销目标", "先盘五项数据，再讲两条路径、三张卡和合作政策。", PALE_GREEN),
        ("政策边界", "竞品数字不能直接当作我方承诺，必须以正式合同为准。", PALE_YELLOW),
    ]
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (title, body, fill) in enumerate(rules):
        c1, c2 = table.rows[i].cells
        c1.width = Cm(3.4)
        c2.width = Cm(12.8)
        for c in (c1, c2):
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(c)
        set_shading(c1, fill)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(title)
        set_run_font(r, 12.3, INK, True)
        p2 = c2.paragraphs[0]
        r = p2.add_run(body)
        set_run_font(r, 11.5, MUTED)

    p = doc.add_paragraph(style="Heading 2")
    p.add_run("内容目录")
    headings = re.findall(r"^## (.+)$", source, flags=re.MULTILINE)
    for i, heading in enumerate(headings, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{i:02d}  ")
        set_run_font(r, 11, ORANGE, True)
        r = p.add_run(heading)
        set_run_font(r, 12.8, INK, True)
    doc.add_page_break()


def add_major_heading(doc: Document, text: str, first: bool):
    if not first:
        doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.6)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_shading(cell, PALE_ORANGE)
    set_cell_border(cell, PALE_ORANGE, 1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    set_run_font(r, 20, ORANGE, True)
    doc.add_paragraph()


def add_table(doc: Document, lines: list[str]):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_shading(cell, GREEN if i == 0 else ("FFFFFF" if i % 2 else PALE_GREY))
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            text = row[j] if j < len(row) else ""
            add_inline_runs(p, text, 10.8 if i else 11.2, "FFFFFF" if i == 0 else INK)
            for run in p.runs:
                run.font.bold = True if i == 0 else run.font.bold
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_quote(doc: Document, text: str):
    fill = PALE_YELLOW if "合规" in text or "必须" in text else PALE_ORANGE
    color = ORANGE if fill == PALE_ORANGE else GREEN
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    set_shading(p, fill)
    set_left_border(p, color)
    add_inline_runs(p, text, 11.8, INK)


def convert(doc: Document, source: str):
    lines = source.splitlines()
    first_major = True
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("> ") and i < 10:
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, table_lines)
            continue

        heading = re.match(r"^(#{2,5})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 2:
                add_major_heading(doc, text, first_major)
                first_major = False
            else:
                style = {3: "Heading 2", 4: "Heading 3", 5: "Heading 4"}.get(level, "Heading 4")
                p = doc.add_paragraph(style=style)
                add_inline_runs(p, text, 17 if level == 3 else 14.7, INK if level == 3 else GREEN)
                p.paragraph_format.keep_with_next = True
            i += 1
            continue

        if stripped.startswith("> "):
            add_quote(doc, stripped[2:])
            i += 1
            continue

        checklist = re.match(r"^- \[([ xX])\]\s+(.+)$", stripped)
        if checklist:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run("☐  " if checklist.group(1) == " " else "☑  ")
            set_run_font(r, 13, GREEN, True)
            add_inline_runs(p, checklist.group(2), 12.8, INK)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.72)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.3
            r = p.add_run("●  ")
            set_run_font(r, 10.5, ORANGE, True)
            add_inline_runs(p, bullet.group(1), 12.8, INK)
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.85)
            p.paragraph_format.first_line_indent = Cm(-0.62)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.3
            r = p.add_run(f"{numbered.group(1)}. ")
            set_run_font(r, 12.8, ORANGE, True)
            add_inline_runs(p, numbered.group(2), 12.8, INK)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.72)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        add_inline_runs(p, stripped, 13.2, INK)
        i += 1


def build():
    source = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "全肌模式招商电销与面销话术｜市场部培训版"
    doc.core_properties.subject = "招商顾问与市场部执行手册"
    doc.core_properties.author = "基于竞品精校直播稿生成"
    add_cover(doc)
    add_quick_rules(doc, source)
    convert(doc, source)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(f"created {build()}")
