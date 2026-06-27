import logging
from docx import Document
from app.core.exceptions import ValidationError

logger = logging.getLogger(__name__)


class DocxService:

    MAX_PARAGRAPHS = 10000

    async def parse(self, file_path: str) -> dict:
        doc = Document(file_path)
        paragraphs = []
        tables = []
        p_count = 0
        t_count = 0

        for element in doc.element.body:
            if element.tag.endswith("}p"):
                if p_count >= self.MAX_PARAGRAPHS:
                    raise ValidationError(f"文档段落数超过 {self.MAX_PARAGRAPHS} 限制")
                p_count += 1
                text = element.text or ""
                for sub in element.iter():
                    if sub.tag.endswith("}t") and sub.text:
                        text += sub.text
                paragraphs.append({
                    "id": f"p{p_count}", "type": "paragraph",
                    "content": text.strip(),
                })
            elif element.tag.endswith("}tbl"):
                t_count += 1
                rows = []
                for tr in element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"):
                    cells = []
                    for tc in tr.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"):
                        cell_text = "".join(
                            t.text or "" for t in tc.iter()
                            if t.tag.endswith("}t") and t.text
                        )
                        cells.append(cell_text.strip())
                    rows.append({"cells": cells})
                tables.append({"id": f"t{t_count}", "type": "table", "rows": rows})

        content = paragraphs + tables
        return {
            "manifest": {
                "file_type": "docx", "version": "1.0.0",
                "paragraph_count": p_count, "table_count": t_count,
            },
            "content": content,
        }

    async def export(self, file_path: str, json_content: dict) -> None:
        doc = Document()
        content = json_content.get("content", json_content) if isinstance(json_content, dict) else json_content

        if isinstance(content, list):
            for item in content:
                if item.get("type") == "paragraph":
                    doc.add_paragraph(item.get("content", ""))
                elif item.get("type") == "table":
                    rows = item.get("rows", [])
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0].get("cells", [])))
                        for i, row_data in enumerate(rows):
                            for j, cell_text in enumerate(row_data.get("cells", [])):
                                if j < len(table.columns):
                                    table.cell(i, j).text = cell_text

        doc.save(file_path)

    def preview_patch(self, json_content: dict, patch: dict) -> dict:
        if patch.get("operation_type") not in ("replace_text", "modify_docx_paragraph"):
            raise ValueError("DOCX 补丁仅支持 replace_text 操作类型")

        content = json_content.get("content", json_content) if isinstance(json_content, dict) else json_content
        target_id = None
        if isinstance(content, list):
            for item in content:
                if item.get("type") == "paragraph" and item.get("content", "").strip():
                    target_id = item["id"]
                    break

        return {
            "preview_passed": True,
            "target_id": target_id,
            "risk_level": "medium",
            "style_loss_risk": True,
        }
