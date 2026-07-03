
"""Sandbox test for docx-parser module."""
import base64
import sys
import tempfile
from pathlib import Path

SAMPLE = Path(__file__).resolve().parent / "samples" / "sample.docx"
MODULE_DIR = Path(__file__).resolve().parents[1]
BACKEND_DIR = MODULE_DIR / "backend"
REPO_BACKEND_DIR = Path(__file__).resolve().parents[3] / "backend"
sys.path.insert(0, str(REPO_BACKEND_DIR))
sys.path.insert(0, str(BACKEND_DIR))

if not SAMPLE.exists():
    print("ERROR: sample.docx not found")
    sys.exit(1)

from docx import Document  # noqa: E402
from docx.shared import Inches  # noqa: E402
from parser import parse_docx_file  # noqa: E402

PNG_1X1_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def validate_shape(result: dict) -> None:
    assert result["format"] == "docx"
    assert isinstance(result["blocks"], list)
    assert isinstance(result["resources"], list)
    assert isinstance(result["resource_diagnostics"], list)
    for block in result["blocks"]:
        assert all(key in block for key in ("type", "text", "page", "resource_ref"))
        assert block["type"] in {"heading", "paragraph", "table", "image"}


def test_sample_docx() -> None:
    result = parse_docx_file(0, SAMPLE)
    validate_shape(result)
    block_types = [block["type"] for block in result["blocks"]]
    texts = [block["text"] for block in result["blocks"]]
    assert block_types[:3] == ["heading", "paragraph", "paragraph"]
    assert "table" in block_types
    assert any("Row1-Col1 | Row1-Col2 | Row1-Col3" in text for text in texts)
    assert result["resources"] == []


def test_generated_image_docx() -> None:
    with tempfile.TemporaryDirectory(prefix="docx-parser-sandbox-") as tmp:
        tmp_dir = Path(tmp)
        image_path = tmp_dir / "pixel.png"
        image_path.write_bytes(base64.b64decode(PNG_1X1_B64))
        docx_path = tmp_dir / "with-image.docx"
        document = Document()
        document.add_heading("Image Case", level=1)
        document.add_paragraph("Before image")
        document.add_picture(str(image_path), width=Inches(0.2))
        document.add_paragraph("After image")
        document.save(str(docx_path))

        result = parse_docx_file(1, docx_path)

    validate_shape(result)
    image_blocks = [block for block in result["blocks"] if block["type"] == "image"]
    assert len(image_blocks) == 1
    assert len(result["resources"]) == 1
    resource = result["resources"][0]
    assert resource["id"] == image_blocks[0]["resource_ref"]
    assert resource["mime_type"] == "image/png"
    assert resource["_bytes_b64"]


def test_empty_docx_returns_empty_blocks() -> None:
    with tempfile.TemporaryDirectory(prefix="docx-parser-sandbox-") as tmp:
        docx_path = Path(tmp) / "empty.docx"
        Document().save(str(docx_path))
        result = parse_docx_file(2, docx_path)

    validate_shape(result)
    assert result["blocks"] == []
    assert result["resources"] == []


def test_corrupt_docx_raises() -> None:
    with tempfile.TemporaryDirectory(prefix="docx-parser-sandbox-") as tmp:
        broken_path = Path(tmp) / "broken.docx"
        broken_path.write_bytes(b"not a docx")
        try:
            parse_docx_file(3, broken_path)
        except Exception:
            return
    raise AssertionError("corrupt DOCX must raise instead of returning fake success")


def main() -> None:
    print("=" * 60)
    print("docx-parser sandbox test")
    print("=" * 60)
    test_sample_docx()
    test_generated_image_docx()
    test_empty_docx_returns_empty_blocks()
    test_corrupt_docx_raises()
    sample = parse_docx_file(0, SAMPLE)
    print("  Validation PASS (%d blocks, %d resources)" % (len(sample["blocks"]), len(sample["resources"])))
    print("PASS: docx-parser sandbox test")


if __name__ == "__main__":
    main()
