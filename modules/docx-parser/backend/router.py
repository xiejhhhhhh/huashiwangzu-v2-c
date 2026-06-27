"""FastAPI router for docx-parser module.

Registers the parse capability with the framework's cross-module registry.
"""
import os
from fastapi import APIRouter, Depends
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import AsyncSessionLocal, get_db
from app.middleware.auth import require_permission
from app.models.user import User
from app.schemas.common import ApiResponse
from app.services.module_registry import register_capability

router = APIRouter(prefix="/api/docx-parser", tags=["docx-parser"])


class ParseRequest(BaseModel):
    file_id: int


def _resolve_user_id(caller: str) -> int:
    from app.core.exceptions import PermissionDenied

    try:
        prefix, raw_id = caller.split(":", 1)
        if prefix == "user":
            return int(raw_id)
    except (TypeError, ValueError):
        pass
    raise PermissionDenied("Invalid caller")


async def _parse(params: dict, caller: str) -> dict:
    """Parse DOCX file into unified content blocks."""
    file_id = int(params.get("file_id", 0))
    if file_id <= 0:
        raise ValueError("file_id must be a positive integer")

    from app.config import get_settings
    from app.core.exceptions import NotFound, ValidationError, AppException
    from app.services.file_service import check_file_access
    from pathlib import Path
    from docx import Document as DocxDocument

    allowed = {"docx"}
    user_id = _resolve_user_id(caller)
    async with AsyncSessionLocal() as db:
        file = await check_file_access(db, file_id, user_id)
        ext = (file.extension or "").lower()
        if ext not in allowed:
            raise ValidationError(f"Unsupported format '{ext}'. Allowed: {', '.join(sorted(allowed))}")
        if not file.storage_path:
            raise NotFound("File storage path is empty")
        upload_root = Path(get_settings().UPLOAD_DIR).resolve()
        full_path = (upload_root / file.storage_path).resolve()
        if os.path.commonpath([str(upload_root), str(full_path)]) != str(upload_root):
            raise AppException("Unsafe file storage path", status_code=400)
        if not full_path.exists() or not full_path.is_file():
            raise NotFound("File on disk not found")

        doc = DocxDocument(str(full_path))
        blocks = []
        resources = []
        resource_counter = 0

        for para in doc.paragraphs:
            text = "\n".join(line.rstrip() for line in para.text.splitlines()).strip()
            if not text:
                continue
            style_name = str(para.style.name) if para.style else ""
            block_type = "标题" if ("heading" in style_name.lower() or "标题" in style_name) else "段落"
            blocks.append({"type": block_type, "text": text, "page": None, "resource_ref": None})

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                blocks.append({"type": "表格", "text": table_text, "page": None, "resource_ref": None})

        for rel in doc.part.rels.values():
            if "image" in str(rel.reltype or "").lower():
                resource_counter += 1
                blocks.append({"type": "图片", "text": "", "page": None, "resource_ref": resource_counter})
                resources.append({
                    "id": resource_counter,
                    "type": "图片",
                    "file_storage_id": None,
                    "text_desc": f"DOCX embedded image ({rel.target_ref})",
                })

    return {
        "file_id": file_id,
        "format": "docx",
        "blocks": blocks,
        "resources": resources,
    }


@router.get("/health")
async def health():
    return ApiResponse(data={"module": "docx-parser", "status": "ok"})


@router.post("/parse")
async def call_parse(payload: ParseRequest, user: User = Depends(require_permission("viewer"))):
    result = await _parse({"file_id": payload.file_id}, f"user:{user.id}")
    return ApiResponse(data=result)


register_capability(
    "docx-parser", "parse", _parse,
    description="Parse DOCX files into unified content blocks",
    brief="解析 DOCX 文档",
    parameters={"file_id": {"type": "int"}},
    min_role="viewer",
)
